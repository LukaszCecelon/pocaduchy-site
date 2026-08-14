"""
Eksport artykulu ze strony do pliku Word.

Powod: kazdy material na pocaduchy.pl ma istniec jako .docx w
D:\\poCADychy_STRONA\\BAZA WIEDZY, w folderze na material. Word jest
powierzchnia do edycji i archiwum, JSON w repozytorium jest tylko formatem
publikacji. Ten skrypt idzie w strone JSON -> Word.

Struktura dokumentu jest ta sama, ktora dostaje Codex przy pisaniu nowych
materialow, zeby oba zrodla dawaly pliki wygladajace tak samo:

  tabela metadanych, Naglowek 1, tresc, Pytania i odpowiedzi,
  Zrodla i status norm, Uwagi dla Claude

Uzycie:
  python scripts/eksport-do-word.py                 # wszystkie artykuly Wiedzy
  python scripts/eksport-do-word.py gwinty-metryczne-tabela
"""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parent.parent
BAZA = Path(r"D:\poCADychy_STRONA\BAZA WIEDZY")

# Nazwa folderu dla artykulu. Luzne dopasowanie do tego, co Lukasz juz ma
# na dysku, zeby nie robic drugiego folderu na ten sam temat.
FOLDERY = {
    "pozycjonowanie-czesci-w-maszynie": "6 SPOSOBÓW POZYCJONOWANIA ELEMENTÓW",
}


def folder_dla(slug, tytul):
    if slug in FOLDERY:
        return FOLDERY[slug]
    return tytul.split(":")[0].strip().upper()


# ---------------------------------------------------------------- markdown

# Inline: **pogrubienie**, `kod`, [tekst](adres). Nie obslugujemy zagniezdzen,
# bo w tresci ich nie ma, a udawana obsluga bylaby gorsza od jawnego braku.
INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def wpisz_inline(akapit, tekst):
    for kawalek in INLINE.split(tekst):
        if not kawalek:
            continue
        if kawalek.startswith("**") and kawalek.endswith("**"):
            akapit.add_run(kawalek[2:-2]).bold = True
        elif kawalek.startswith("`") and kawalek.endswith("`"):
            bieg = akapit.add_run(kawalek[1:-1])
            bieg.font.name = "Consolas"
        elif kawalek.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", kawalek)
            bieg = akapit.add_run(m.group(1))
            bieg.underline = True
            bieg.font.color.rgb = RGBColor(0x96, 0x45, 0x3B)
            # Adres zostaje widoczny, bo Lukasz musi wiedziec, dokad link
            # prowadzi, gdy bedzie tekst przestawial.
            akapit.add_run(f" ({m.group(2)})").font.size = Pt(8)
        else:
            akapit.add_run(kawalek)


def tabela_markdown(dok, linie):
    wiersze = [
        [k.strip() for k in w.strip().strip("|").split("|")]
        for w in linie
        if not re.match(r"^\s*\|[\s:|-]+\|\s*$", w)
    ]
    if not wiersze:
        return
    t = dok.add_table(rows=len(wiersze), cols=len(wiersze[0]))
    t.style = "Table Grid"
    for i, wiersz in enumerate(wiersze):
        for j, komorka in enumerate(wiersz[: len(wiersze[0])]):
            cel = t.cell(i, j)
            cel.text = ""
            p = cel.paragraphs[0]
            wpisz_inline(p, komorka)
            if i == 0:
                for bieg in p.runs:
                    bieg.bold = True


def wpisz_markdown(dok, tekst):
    linie = tekst.split("\n")
    i = 0
    while i < len(linie):
        linia = linie[i]

        if not linia.strip():
            i += 1
            continue

        if linia.lstrip().startswith("|"):
            blok = []
            while i < len(linie) and linie[i].lstrip().startswith("|"):
                blok.append(linie[i])
                i += 1
            tabela_markdown(dok, blok)
            continue

        m = re.match(r"^(#{2,4})\s+(.*)", linia)
        if m:
            poziom = min(len(m.group(1)), 4)
            dok.add_paragraph(m.group(2).strip(), style=f"Heading {poziom}")
            i += 1
            continue

        m = re.match(r"^\s*[-*]\s+(.*)", linia)
        if m:
            wpisz_inline(dok.add_paragraph(style="List Bullet"), m.group(1))
            i += 1
            continue

        m = re.match(r"^\s*\d+\.\s+(.*)", linia)
        if m:
            wpisz_inline(dok.add_paragraph(style="List Number"), m.group(1))
            i += 1
            continue

        wpisz_inline(dok.add_paragraph(), linia.strip())
        i += 1


# ------------------------------------------------------------------ bloki

# Tabele liczone z danych. W Wordzie nie ma sensu ich odtwarzac, bo strona
# i tak generuje je z modulu w src/lib. Zostawiamy jawny znacznik, zeby
# Lukasz wiedzial, ze w tym miejscu cos bedzie, i nie przepisywal tego recznie.
GENEROWANE = {
    "tabelaPierscieni": "tablica rowkow pod pierscienie osadcze, liczona z danych",
    "tabelaGwintow": "tablica gwintow metrycznych, 43 wiersze, liczona z danych",
    "tabelaChropowatosci": "tablice Ra i Rz dla 29 metod obrobki, liczone z danych",
}


def znacznik(dok, tekst):
    p = dok.add_paragraph()
    bieg = p.add_run(tekst)
    bieg.bold = True
    bieg.font.color.rgb = RGBColor(0x96, 0x45, 0x3B)
    return p


def wpisz_bloki(dok, bloki):
    for b in bloki:
        typ = b.get("type")

        if typ == "tekst":
            wpisz_markdown(dok, b.get("body", ""))

        elif typ == "tabela":
            wpisz_markdown(dok, b.get("markdown", ""))

        elif typ in GENEROWANE:
            znacznik(dok, f"[TABELA GENEROWANA] {GENEROWANE[typ]}. Nie edytuj tutaj.")
            if b.get("podpis"):
                dok.add_paragraph(f"Podpis: {b['podpis']}")

        elif typ == "rysunek":
            znacznik(dok, "[SVG] Rysunek schematyczny, rysowany kodem. Nic nie robisz.")
            if b.get("podpis"):
                dok.add_paragraph(f"Podpis: {b['podpis']}")

        elif typ == "obraz":
            znacznik(dok, f"[OBRAZ] Na stronie: {b.get('src', '')}")
            if b.get("podpis"):
                dok.add_paragraph(f"Podpis: {b['podpis']}")

        elif typ == "wzor":
            p = dok.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            bieg = p.add_run(b.get("latex", ""))
            bieg.font.name = "Consolas"

        elif typ == "wideo":
            znacznik(dok, f"[WIDEO] Na stronie: {b.get('src', '')}")

        elif typ == "plik":
            znacznik(dok, f"[PLIK DO POBRANIA] {b.get('tytul', '')}")


# ---------------------------------------------------------------- dokument


def zbuduj(dane, sciezka):
    dok = Document()

    tabela = dok.add_table(rows=8, cols=2)
    tabela.style = "Table Grid"
    pola = [
        ("slug", dane.get("slug", "")),
        ("kategoria", dane.get("kategoria", "")),
        ("tytuł", dane.get("title", "")),
        ("seoTitle", dane.get("seoTitle", "")),
        ("description", dane.get("description", "")),
        ("słowa kluczowe", ", ".join(dane.get("slowaKluczowe", []))),
        ("data", dane.get("date", "")),
        ("status", "opublikowane, ten plik to kopia do edycji"),
    ]
    for i, (nazwa, wartosc) in enumerate(pola):
        tabela.cell(i, 0).text = nazwa
        for bieg in tabela.cell(i, 0).paragraphs[0].runs:
            bieg.bold = True
        tabela.cell(i, 1).text = str(wartosc)

    p = dok.add_paragraph()
    bieg = p.add_run(
        "Ten materiał jest już na stronie. Popraw tekst, powstawiaj obrazy "
        "i odeślij plik. Zmiany przeniosę na stronę i opublikuję. Nie zmieniaj "
        "pola slug, bo to adres strony."
    )
    bieg.italic = True

    dok.add_paragraph(dane.get("title", ""), style="Heading 1")

    wpisz_bloki(dok, dane.get("blocks", []))

    if dane.get("faq"):
        dok.add_paragraph("Pytania i odpowiedzi", style="Heading 2")
        for q in dane["faq"]:
            dok.add_paragraph(q.get("pytanie", ""), style="Heading 3")
            wpisz_inline(dok.add_paragraph(), q.get("odpowiedz", ""))

    dok.add_paragraph("Uwagi dla Claude", style="Heading 2")
    p = dok.add_paragraph()
    p.add_run(
        "Tu wpisz, co mam zmienić na stronie, czego brakuje albo co przenieść "
        "w inne miejsce. Skasuj tę sekcję, jeśli nie masz uwag."
    ).italic = True

    sciezka.parent.mkdir(parents=True, exist_ok=True)
    dok.save(sciezka)
    return sciezka


def main():
    wybrane = sys.argv[1:]
    katalog = REPO / "content" / "wiedza"
    zrobione = []

    for plik in sorted(katalog.glob("*.json")):
        dane = json.loads(plik.read_text(encoding="utf-8"))
        if wybrane and dane["slug"] not in wybrane:
            continue
        folder = BAZA / folder_dla(dane["slug"], dane["title"])
        sciezka = folder / f"{dane['slug']}.docx"
        zbuduj(dane, sciezka)
        zrobione.append(sciezka)

    for s in zrobione:
        print(f"{s.stat().st_size:>8} B  {s}")
    print(f"[eksport-do-word] OK: {len(zrobione)} plikow")


if __name__ == "__main__":
    main()
