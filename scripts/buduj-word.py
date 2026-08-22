"""
Sklada dokument Word do bazy wiedzy z prostego opisu w JSON.

Format wyjsciowy jest dokladnie ten sam, ktory powstawal wczesniej maszynowo:
tabela metadanych, jeden Naglowek 1, tresc, Pytania i odpowiedzi,
Zrodla i status norm, Uwagi dla Claude. Dzieki temu wszystkie pliki w
D:\\poCADychy_STRONA\\BAZA WIEDZY wygladaja tak samo niezaleznie od tego,
kto je napisal, i daja sie tak samo sprawdzac oraz przenosic na strone.

Uzycie:
  python scripts/buduj-word.py opis.json
  python scripts/buduj-word.py katalog-z-opisami/

Format opisu: patrz scripts/buduj-word-przyklad.json
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

BAZA = Path(r"D:\poCADychy_STRONA\BAZA WIEDZY\SEKCJA-RYSUNEK TECHNICZNY")
RDZA = RGBColor(0x96, 0x45, 0x3B)

POLA = ["slug", "kategoria", "tytuł", "seoTitle", "description",
        "słowa kluczowe", "data", "status"]


def akapit_wyrozniony(dok, tekst):
    p = dok.add_paragraph()
    bieg = p.add_run(tekst)
    bieg.bold = True
    bieg.font.color.rgb = RDZA
    return p


def tabela(dok, wiersze):
    t = dok.add_table(rows=len(wiersze), cols=len(wiersze[0]))
    t.style = "Table Grid"
    for i, wiersz in enumerate(wiersze):
        for j, kom in enumerate(wiersz):
            cel = t.cell(i, j)
            cel.text = ""
            p = cel.paragraphs[0]
            bieg = p.add_run(str(kom))
            if i == 0:
                bieg.bold = True
    return t


def blok(dok, b):
    typ = b.get("typ")

    if typ == "h2":
        dok.add_paragraph(b["tekst"], style="Heading 2")
    elif typ == "h3":
        dok.add_paragraph(b["tekst"], style="Heading 3")
    elif typ == "p":
        dok.add_paragraph(b["tekst"])
    elif typ == "lista":
        for x in b["pozycje"]:
            dok.add_paragraph(x, style="List Bullet")
    elif typ == "kroki":
        for x in b["pozycje"]:
            dok.add_paragraph(x, style="List Number")
    elif typ == "tabela":
        tabela(dok, b["wiersze"])
    elif typ in ("svg", "inventor", "zdjecie"):
        znacznik = {"svg": "[SVG]", "inventor": "[INVENTOR]", "zdjecie": "[ZDJĘCIE]"}[typ]
        akapit_wyrozniony(dok, f"{znacznik} {b['opis']}")
        dok.add_paragraph(f"Podpis: {b['podpis']}")
    else:
        raise ValueError(f"nieznany typ bloku: {typ}")


def zbuduj(opis):
    dok = Document()

    meta = opis["meta"]
    wartosci = [
        ("slug", meta["slug"]),
        ("kategoria", "rysunek-techniczny"),
        ("tytuł", meta["tytul"]),
        ("seoTitle", meta["seoTitle"]),
        ("description", meta["description"]),
        ("słowa kluczowe", ", ".join(meta["slowaKluczowe"])),
        ("data", meta.get("data", "2026-08-14")),
        ("status", "do edycji"),
    ]
    t = dok.add_table(rows=len(wartosci), cols=2)
    t.style = "Table Grid"
    for i, (nazwa, w) in enumerate(wartosci):
        t.cell(i, 0).text = ""
        t.cell(i, 0).paragraphs[0].add_run(nazwa).bold = True
        t.cell(i, 1).text = str(w)

    p = dok.add_paragraph()
    p.add_run(
        # Bez nawiasow kwadratowych, bo skrypt kontrolny liczy znaczniki
        # w calym tekscie i to zdanie zawyzaloby wynik.
        "Popraw tekst, powstawiaj obrazy w miejscach oznaczonych znacznikiem "
        "INVENTOR albo ZDJĘCIE, potem odeślij plik. Bez odesłania nic nie "
        "trafia na stronę. "
        "Nie zmieniaj pola slug, bo to przyszły adres strony."
    ).italic = True

    dok.add_paragraph(meta["tytul"], style="Heading 1")

    for b in opis["tresc"]:
        blok(dok, b)

    dok.add_paragraph("Pytania i odpowiedzi", style="Heading 2")
    for q in opis["faq"]:
        dok.add_paragraph(q["pytanie"], style="Heading 3")
        dok.add_paragraph(q["odpowiedz"])

    dok.add_paragraph("Źródła i status norm", style="Heading 2")
    for z in opis["zrodla"]:
        dok.add_paragraph(z)
    if opis.get("niepotwierdzone"):
        dok.add_paragraph("Nie potwierdzone bezpośrednio:")
        for z in opis["niepotwierdzone"]:
            dok.add_paragraph(z, style="List Bullet")

    dok.add_paragraph("Uwagi dla Claude", style="Heading 2")
    for u in opis.get("uwagi", ["Brak uwag."]):
        dok.add_paragraph(u)

    katalog = BAZA / f"{opis['numer']:02d}-{meta['slug']}"
    katalog.mkdir(parents=True, exist_ok=True)
    sciezka = katalog / f"{opis['numer']:02d}-{meta['slug']}.docx"
    dok.save(sciezka)
    return sciezka


def main():
    cel = Path(sys.argv[1])
    pliki = sorted(cel.glob("*.json")) if cel.is_dir() else [cel]
    for p in pliki:
        opis = json.loads(p.read_text(encoding="utf-8"))
        s = zbuduj(opis)
        print(f"{s.stat().st_size:>7} B  {s.name}")


if __name__ == "__main__":
    main()
