"""
Eksport karty odcinka do pliku Word do przegladu przez Lukasza.

Odcinek ma inna strukture niz artykul Wiedzy: zamiast listy blokow ma lead,
liste "czego sie dowiesz", sekcje z akapitami i wnioski. Dlatego osobny
skrypt, a nie doklejanie warunkow do eksport-do-word.py.

Uzycie: python scripts/odcinek-do-word.py content/_poczekalnia/odcinek-*.json
"""

import json
import sys
from pathlib import Path

from docx import Document

BAZA = Path(r"D:\poCADychy_STRONA\BAZA WIEDZY")


def zbuduj(dane, nazwa):
    d = Document()

    pola = [
        ("videoId", dane["videoId"]),
        ("tytuł", dane["title"]),
        ("seoTitle", dane["seoTitle"]),
        ("description", dane["description"]),
        ("czas trwania", dane["czasTrwania"]),
        ("status", "do akceptacji, nieopublikowane"),
    ]
    t = d.add_table(rows=len(pola), cols=2)
    t.style = "Table Grid"
    for i, (k, v) in enumerate(pola):
        t.cell(i, 0).text = ""
        t.cell(i, 0).paragraphs[0].add_run(k).bold = True
        t.cell(i, 1).text = str(v)

    p = d.add_paragraph()
    p.add_run(
        "Karta odcinka przygotowana z transkrypcji. Popraw, co uznasz za "
        "stosowne, i odeślij. Bez odesłania nie trafia na stronę. Pole "
        "videoId muszę uzupełnić przed publikacją."
    ).italic = True

    d.add_paragraph(dane["title"], style="Heading 1")
    d.add_paragraph(dane["lead"])

    d.add_paragraph("Czego dotyczy ten odcinek", style="Heading 2")
    for x in dane["czegoSieDowiesz"]:
        d.add_paragraph(x, style="List Bullet")

    for s in dane["sekcje"]:
        d.add_paragraph(s["tytul"], style="Heading 2")
        for a in s["akapity"]:
            d.add_paragraph(a)

    d.add_paragraph("Co z tego wynika", style="Heading 2")
    for w in dane["wnioski"]:
        d.add_paragraph(w, style="List Bullet")

    d.add_paragraph("Uwagi dla Claude", style="Heading 2")
    for u in dane.get("_uwagi", ["Brak uwag."]):
        d.add_paragraph(u)

    katalog = BAZA / nazwa
    katalog.mkdir(parents=True, exist_ok=True)
    sciezka = katalog / f"{Path(sys.argv[1]).stem}.docx"
    d.save(sciezka)
    return sciezka


if __name__ == "__main__":
    dane = json.loads(Path(sys.argv[1]).read_text(encoding="utf-8"))
    s = zbuduj(dane, "ODCINEK 16 GWINTY")
    print(f"{s.stat().st_size} B  {s}")
